import logging
import os
import uuid
from datetime import datetime
from database import db_instance
from docx import Document
from docx.shared import Inches
import platform

logger = logging.getLogger(__name__)

async def generate_report(inspection_id: str, vessel_id: str = None, visit_id: str = None):
    logger.info(f"[{inspection_id}] Generating Final Reports...")
    
    session = await db_instance.analysis_sessions.find_one({"session_id": inspection_id})
    defects = await db_instance.defect_registry.find({"inspection_id": inspection_id}).to_list(length=None)
    costs = await db_instance.repair_estimations.find({"inspection_id": inspection_id}).to_list(length=None)
    
    reports_dir = os.path.join("outputs", "sessions", inspection_id, "reports")
    os.makedirs(reports_dir, exist_ok=True)
    
    doc = Document()
    doc.add_heading('Maritime Inspection Report', 0)
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(f"Inspection ID: {inspection_id}")
    doc.add_paragraph(f"Date Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    
    if session:
        doc.add_paragraph(f"Vessel ID: {vessel_id}")
        doc.add_paragraph(f"Visit ID: {visit_id}")
        
    doc.add_heading('Defect Summary', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Defect Type'
    hdr_cells[1].text = 'Part'
    hdr_cells[2].text = 'Severity'
    hdr_cells[3].text = 'Estimated Cost'
    
    cost_map = {c["defect_id"]: c for c in costs}
    total_cost = 0.0
    
    for defect in defects:
        row_cells = table.add_row().cells
        row_cells[0].text = defect.get("defect_type", "Unknown")
        row_cells[1].text = defect.get("part_name", "Unknown")
        row_cells[2].text = defect.get("severity", "LOW")
        
        c_est = cost_map.get(defect["defect_id"])
        if c_est:
            cost_val = c_est.get("estimated_cost", 0.0)
            total_cost += cost_val
            row_cells[3].text = f"${cost_val:,.2f}"
        else:
            row_cells[3].text = "N/A"
            
    doc.add_paragraph(f"\nTotal Estimated Repair Exposure: ${total_cost:,.2f}")
    
    doc.add_heading('Defect Images', level=1)
    for defect in defects:
        crop_path = defect.get("crop_path")
        if crop_path and os.path.exists(crop_path):
            doc.add_paragraph(f"{defect.get('defect_type')} on {defect.get('part_name')} (Severity: {defect.get('severity')})")
            doc.add_picture(crop_path, width=Inches(4.0))
            
    docx_path = os.path.join(reports_dir, "report.docx")
    pdf_path = os.path.join(reports_dir, "report.pdf")
    
    doc.save(docx_path)
    logger.info(f"[{inspection_id}] Saved DOCX report to {docx_path}")
    
    # PDF Conversion (if docx2pdf works on this OS environment, typically requires Windows/macOS with Word installed)
    # Since docx2pdf uses COM on Windows, it might fail if Word isn't installed.
    # We will try/except it.
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        logger.info(f"[{inspection_id}] Saved PDF report to {pdf_path}")
    except Exception as e:
        logger.warning(f"[{inspection_id}] PDF conversion failed or not supported in this environment: {e}. Generating empty PDF placeholder.")
        # Fallback empty PDF if we can't generate
        with open(pdf_path, "wb") as f:
            f.write(b"%PDF-1.4\n%Placeholder PDF\n")
            
    # Insert to DB
    from models import ReportRecord
    records = []
    for path in [docx_path, pdf_path]:
        r = ReportRecord(
            report_id=str(uuid.uuid4()),
            inspection_id=inspection_id,
            vessel_id=vessel_id,
            visit_id=visit_id,
            file_path=os.path.abspath(path)
        )
        records.append(r.model_dump())
        
    if records:
        await db_instance.reports.insert_many(records)
        logger.info(f"[{inspection_id}] Inserted {len(records)} report records into MongoDB.")
